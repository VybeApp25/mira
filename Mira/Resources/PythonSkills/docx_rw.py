#!/usr/bin/env python3
"""Read or write a .docx file.
Args JSON (stdin): {op: "read"|"write"|"append", path, content?, heading?}
  read  — returns all paragraph text
  write — creates/overwrites the file with content (markdown-like headings supported)
  append — appends a paragraph (with optional heading before it)
Output JSON (stdout): {success, text?, path?, error?}
"""
import sys, json

def main():
    try:
        args = json.load(sys.stdin)
    except Exception:
        args = {}
    op      = args.get("op", "read")
    path    = args.get("path", "")
    content = args.get("content", "")
    heading = args.get("heading")

    try:
        from docx import Document
    except ImportError:
        print(json.dumps({"success": False, "error": "python-docx not installed. Run: pip3 install python-docx"}))
        return

    try:
        if op == "read":
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            print(json.dumps({"success": True, "text": text, "paragraphs": len(doc.paragraphs)}))

        elif op == "write":
            doc = Document()
            for line in content.splitlines():
                stripped = line.strip()
                if stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped:
                    doc.add_paragraph(stripped)
            doc.save(path)
            print(json.dumps({"success": True, "path": path}))

        elif op == "append":
            doc = Document(path)
            if heading:
                doc.add_heading(heading, level=2)
            if content:
                doc.add_paragraph(content)
            doc.save(path)
            print(json.dumps({"success": True, "path": path}))

        else:
            print(json.dumps({"success": False, "error": f"Unknown op: {op}"}))

    except Exception as e:
        print(json.dumps({"success": False, "error": str(e)}))

main()
